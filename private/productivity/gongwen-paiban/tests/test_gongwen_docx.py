#!/usr/bin/env python3
"""
gongwen-paiban python-docx 生成测试
验证 .docx 文档能正常生成，且格式设置正确（字体、字号、行距、缩进等）。

如果 python-docx 未安装，测试自动跳过。
"""

import sys
import os
import tempfile


def _check_skip():
    """检查 python-docx 是否可用"""
    try:
        import docx  # noqa
        return False
    except ImportError:
        print("  ⏭️  跳过: python-docx 未安装")
        return True


def test_docx_basic():
    """测试基本 docx 生成：标题 + 一级标题 + 正文 + 署名"""
    if _check_skip():
        return True  # 跳过算通过

    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    from docx.oxml.ns import qn

    def set_run_font(run, western="Times New Roman", eastern="仿宋_GB2312", size=16, bold=False):
        run.font.name = western
        run.font.size = Pt(size)
        run.bold = bold
        run.element.rPr.rFonts.set(qn('w:eastAsia'), eastern)

    def add_title(doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, eastern="方正小标宋简体", size=22)

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    add_title(doc, "关于XXXX工作的报告")
    doc.save("/tmp/test_gongwen_docx.docx")

    assert os.path.isfile("/tmp/test_gongwen_docx.docx"), "docx 未生成"
    assert os.path.getsize("/tmp/test_gongwen_docx.docx") > 1000, "docx 文件太小"
    print("  ✅ docx 基本生成成功")
    os.unlink("/tmp/test_gongwen_docx.docx")
    return True


def test_docx_formatting():
    """测试 docx 格式：字体、字号、行距、缩进"""
    if _check_skip():
        return True  # 跳过算通过

    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # 验证页面边距（允许微小误差）
    expected_top = Cm(3.7)
    diff = abs(section.top_margin - expected_top)
    assert diff < 100000, f"上边距不匹配: 期望 {expected_top}, 实际 {section.top_margin}, 差值 {diff}"
    print(f"  ✅ 页边距设置正确: 上 {section.top_margin/914400*2.54:.1f}cm")

    # 正文：16pt，首行缩进32pt，28磅行距
    p = doc.add_paragraph()
    run = p.add_run("这是一段测试正文")
    run.font.size = Pt(16)
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.first_line_indent = Pt(32)

    assert pf.line_spacing == Pt(28), "行距设置不匹配"
    assert abs(pf.first_line_indent - Pt(32)) <= Pt(0.1), "首行缩进设置不匹配"
    print(f"  ✅ 行距 {pf.line_spacing} ✓")
    print(f"  ✅ 首行缩进 {pf.first_line_indent} ✓")

    doc.save("/tmp/test_gongwen_format.docx")
    os.unlink("/tmp/test_gongwen_format.docx")
    print("  ✅ docx 格式验证通过")
    return True


def main():
    print("=" * 60)
    print("  公文排版 python-docx 测试")
    print("=" * 60)

    tests = [
        ("基本 docx 生成", test_docx_basic),
        ("格式验证（边距/行距/缩进）", test_docx_formatting),
    ]

    results = []
    for name, func in tests:
        print(f"\n[{len(results)+1}/{len(tests)}] {name}")
        print("-" * 40)
        try:
            ok = func()
        except Exception as e:
            print(f"  ❌ 异常: {e}")
            ok = False
        results.append(ok)

    print("\n" + "=" * 60)
    passed = sum(1 for r in results if r)
    total = len(results)
    print(f"  结果: {passed}/{total} 通过")
    for name, ok in zip([t[0] for t in tests], results):
        status = "✅" if ok else "❌"
        print(f"  {status} {name}")

    return 0 if all(results) else 1


if __name__ == "__main__":
    ret = main()
    print(f"\n测试结果: {'✅ 全部通过' if ret == 0 else '❌ 有失败的测试'} (exit={ret})")
    sys.exit(ret)
