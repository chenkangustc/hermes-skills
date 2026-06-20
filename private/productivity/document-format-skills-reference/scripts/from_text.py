#!/usr/bin/env python3
"""Create a DOCX from plain text or Markdown, then optionally format it."""

import argparse
import logging
import os
import re
import sys
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT_DIR = Path(__file__).resolve().parent.parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from scripts.formatter import (  # noqa: E402
    PRESETS,
    _merge_preset_settings,
    format_document,
    load_custom_preset,
    load_preset_file,
)
from scripts.punctuation import process_document as fix_punctuation  # noqa: E402


logger = logging.getLogger("docformat.from_text")


def detect_markdown(text):
    """Heuristically detect whether text looks like Markdown."""
    if not text or not text.strip():
        return False
    text = text.lstrip("\ufeff")

    score = 0
    lines = text.split("\n")

    if sum(1 for line in lines if re.match(r"^\s*#{1,6}\s+\S", line)) >= 1:
        score += 3
    if len(re.findall(r"\*\*[^*\n]+\*\*", text)) >= 2:
        score += 2
    if sum(1 for line in lines if re.match(r"^\s*[-*+]\s+\S", line)) >= 2:
        score += 2
    if "```" in text:
        score += 2
    if sum(1 for line in lines if line.startswith("> ")) >= 1:
        score += 1

    return score >= 3


def parse_markdown_inline(text):
    """Return [(text, is_bold), ...] for **bold** and __bold__ spans."""
    parts = []
    pattern = re.compile(r"\*\*([^*\n]+)\*\*|__([^_\n]+)__")
    last_end = 0

    for match in pattern.finditer(text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        bold_text = match.group(1) or match.group(2)
        if bold_text:
            parts.append((bold_text, True))
        last_end = match.end()

    if last_end < len(text):
        parts.append((text[last_end:], False))

    return parts or [(text, False)]


def _add_title(doc, title):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title.strip())
    run.font.size = Pt(22)
    return para


def create_docx_from_text(title, body_text, output_path):
    """Create a minimal DOCX from title and plain text."""
    doc = Document()
    _add_title(doc, title)

    body_text = body_text.strip()
    if body_text:
        for block in body_text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            for line in block.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

    doc.save(output_path)
    return output_path


def create_docx_from_markdown(title, md_text, output_path):
    """Create a DOCX from a small, document-oriented Markdown subset."""
    doc = Document()
    lines = md_text.split("\n")

    cn_nums = [
        "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
        "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
    ]

    def cn(n):
        return cn_nums[n - 1] if 1 <= n <= len(cn_nums) else str(n)

    h2_counter = 0
    h3_counters = {}
    h4_counters = {}
    current_h2 = 0
    current_h3 = 0
    title_added = False
    in_code_block = False
    code_buffer = []

    def add_para_with_inline(text, alignment=None, size=None):
        para = doc.add_paragraph()
        if alignment is not None:
            para.alignment = alignment
        for content, is_bold in parse_markdown_inline(text):
            if not content:
                continue
            run = para.add_run(content)
            run.font.bold = is_bold
            if size:
                run.font.size = Pt(size)
        return para

    def flush_code():
        for code_line in code_buffer:
            if code_line.strip():
                doc.add_paragraph(code_line)
        code_buffer.clear()

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not line.strip():
            continue

        heading = re.match(r"^\s*(#{1,6})\s*(.*)$", line)
        if heading:
            level = len(heading.group(1))
            content = heading.group(2).strip()
            if level == 1 and not title_added:
                add_para_with_inline(content, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=22)
                title_added = True
            elif level <= 2:
                h2_counter += 1
                current_h2 = h2_counter
                add_para_with_inline(f"{cn(h2_counter)}、{content}")
            elif level == 3:
                h3_counters[current_h2] = h3_counters.get(current_h2, 0) + 1
                current_h3 = h3_counters[current_h2]
                add_para_with_inline(f"（{cn(current_h3)}）{content}")
            else:
                key = (current_h2, current_h3)
                h4_counters[key] = h4_counters.get(key, 0) + 1
                add_para_with_inline(f"{h4_counters[key]}. {content}")
            continue

        unordered = re.match(r"^\s*[-*+]\s+(.*)$", line)
        ordered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        quote = re.match(r"^>\s+(.*)$", line)
        if unordered:
            add_para_with_inline(unordered.group(1).strip())
        elif ordered:
            add_para_with_inline(ordered.group(1).strip())
        elif quote:
            add_para_with_inline(quote.group(1).strip())
        else:
            add_para_with_inline(line.strip())

    if in_code_block:
        flush_code()

    if not title_added and title.strip():
        title_para = _add_title(doc, title)
        body = doc.element.body
        body.remove(title_para._p)
        body.insert(0, title_para._p)

    doc.save(output_path)
    return output_path


def _format_overrides(args):
    overrides = {}
    if args.deep_clean:
        overrides["deep_clean"] = True
    if args.smart_table_align:
        overrides.setdefault("table", {})["smart_align"] = True
    if args.no_page_number:
        overrides["page_number"] = False
    return overrides


def _load_custom_settings(args):
    settings = load_preset_file(args.custom_settings) if args.custom_settings else None
    overrides = _format_overrides(args)
    if args.preset == "custom" and overrides and settings is None:
        settings = load_custom_preset() or deepcopy(PRESETS["official"])
    if overrides:
        settings = _merge_preset_settings(settings or {}, overrides)
    return settings


def generate_and_process(args):
    text = Path(args.input).read_text(encoding=args.encoding).lstrip("\ufeff")
    title = args.title or Path(args.input).stem
    use_markdown = args.markdown or (not args.plain and detect_markdown(text))

    if args.no_process:
        raw_output = args.output
    else:
        fd, raw_output = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

    try:
        if use_markdown:
            create_docx_from_markdown(title, text, raw_output)
            logger.info("Generated raw DOCX from Markdown")
        else:
            create_docx_from_text(title, text, raw_output)
            logger.info("Generated raw DOCX from plain text")

        if args.no_process:
            logger.info("Saved: %s", args.output)
            return args.output

        fd, fixed_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            fix_punctuation(raw_output, fixed_path, space_mode=args.space_mode)
            format_document(
                fixed_path,
                args.output,
                preset_name=args.preset,
                revision_mode=args.revision,
                bold_serial=not args.no_bold_serial,
                custom_settings=_load_custom_settings(args),
            )
        finally:
            if os.path.exists(fixed_path):
                os.unlink(fixed_path)

        logger.info("Saved: %s", args.output)
        return args.output
    finally:
        if not args.no_process and os.path.exists(raw_output):
            os.unlink(raw_output)


def main(argv=None):
    parser = argparse.ArgumentParser(
        description="Create a DOCX from .txt/.md content and optionally apply smart formatting."
    )
    parser.add_argument("input", help="Input text or Markdown file")
    parser.add_argument("output", help="Output .docx file")
    parser.add_argument("--title", help="Document title; defaults to input file stem")
    parser.add_argument("--encoding", default="utf-8-sig", help="Input file encoding")
    parser.add_argument("--markdown", action="store_true", help="Treat input as Markdown")
    parser.add_argument("--plain", action="store_true", help="Treat input as plain text")
    parser.add_argument("--no-process", action="store_true", help="Only generate DOCX, without punctuation/formatting")
    parser.add_argument(
        "--space-mode",
        choices=("remove_all", "keep_en_boundary", "keep_all"),
        default="remove_all",
        help="Spacing mode used during punctuation processing.",
    )
    parser.add_argument(
        "--preset",
        choices=tuple(PRESETS.keys()) + ("custom",),
        default="official",
        help="Formatting preset.",
    )
    parser.add_argument("--custom-settings", help="Custom preset/config JSON file")
    parser.add_argument("--revision", action="store_true", help="Output supported formatting changes as revisions")
    parser.add_argument("--no-bold-serial", action="store_true", help="Do not bold 一是/二是 body prefixes")
    parser.add_argument("--deep-clean", action="store_true", help="Clear inherited formatting before applying preset")
    parser.add_argument("--smart-table-align", action="store_true", help="Align table cells by content type")
    parser.add_argument("--no-page-number", action="store_true", help="Skip adding page numbers")
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging level.",
    )
    args = parser.parse_args(argv)

    if args.markdown and args.plain:
        parser.error("--markdown and --plain cannot be used together")

    logging.basicConfig(level=getattr(logging, args.log_level), format="%(message)s")
    generate_and_process(args)


if __name__ == "__main__":
    main()
