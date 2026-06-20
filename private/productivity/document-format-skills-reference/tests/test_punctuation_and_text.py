import tempfile
from pathlib import Path

from docx import Document

from scripts.from_text import create_docx_from_markdown, detect_markdown
from scripts.punctuation import _process_spaces_text, fix_text, process_paragraph


def test_punctuation_protects_special_patterns():
    text = "会议时间:上午9:30至下午14:30,请发送至 report@gov.cn,参照 ISO 9001:2015 执行."
    fixed = fix_text(text)
    assert "9:30" in fixed
    assert "14:30" in fixed
    assert "report@gov.cn" in fixed
    assert "ISO 9001:2015" in fixed
    assert "会议时间：" in fixed
    assert "执行。" in fixed


def test_space_modes():
    assert _process_spaces_text("中 文 A测 试 1项", "remove_all") == "中文A测试1项"
    assert _process_spaces_text("中 文 A测 试 1项", "keep_en_boundary") == "中文 A 测试 1 项"
    assert _process_spaces_text("中 文 A测 试 1项", "keep_all") == "中 文 A测 试 1项"


def test_process_paragraph_preserves_run_count_for_same_length_changes():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("他说")
    paragraph.add_run('"好"')

    assert process_paragraph(paragraph, space_mode="keep_all")
    assert paragraph.text == "他说“好”"
    assert len(paragraph.runs) == 2


def test_markdown_generation_maps_headings():
    markdown = """# 主标题

## 工作背景
正文内容。

### 总体情况
**重点**内容。
"""
    assert detect_markdown(markdown)
    assert detect_markdown("\ufeff" + markdown.lstrip())

    with tempfile.TemporaryDirectory() as folder:
        output = Path(folder) / "out.docx"
        create_docx_from_markdown("备用标题", markdown, output)
        doc = Document(output)

    texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    assert texts[0] == "主标题"
    assert texts[1] == "一、工作背景"
    assert texts[3] == "（一）总体情况"
    assert texts[4] == "重点内容。"
    assert doc.paragraphs[4].runs[0].font.bold is True
