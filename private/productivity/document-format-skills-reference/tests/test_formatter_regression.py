import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from scripts.formatter import PRESETS, add_page_number, detect_para_type, format_document


def _detect_chain(texts, doc_idx_offset=0):
    results = []
    previous = None
    for i, text in enumerate(texts):
        para_type = detect_para_type(
            text=text,
            index=i + doc_idx_offset,
            total=len(texts) + doc_idx_offset,
            alignment=None,
            all_texts=texts,
            all_texts_index=i,
            prev_para_type=previous,
        )
        results.append(para_type)
        previous = para_type
    return results


def _footer_text(footer):
    return "".join(paragraph.text for paragraph in footer.paragraphs)


def test_multiline_title_and_recipient_detection():
    texts = [
        "某某市人民政府办公室",
        "关于印发2026年度重点工作任务清单的",
        "通知",
        "各部门、各单位：",
        "现将有关事项通知如下。",
    ]
    assert _detect_chain(texts) == ["title", "title", "title", "recipient", "body"]


def test_dotted_date_not_misidentified_as_heading():
    texts = [f"第{i}段正文。" for i in range(20)] + [
        "某某市发展和改革委员会",
        "2026.04.20",
    ]
    assert _detect_chain(texts)[-2:] == ["signature", "date"]


def test_official_page_number_outside_spacing_and_offset():
    doc = Document()
    section = doc.sections[0]
    section.bottom_margin = Cm(3.5)

    add_page_number(doc, style="dash", position="outside", offset_from_text_mm=7)

    assert abs(section.footer_distance.cm - 2.8) < 0.02
    assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert section.even_page_footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert _footer_text(section.footer).endswith(" —　")
    assert _footer_text(section.even_page_footer).startswith("　— ")


def test_page_number_styles_and_replacement():
    with tempfile.TemporaryDirectory() as folder:
        source = Path(folder) / "source.docx"
        output = Path(folder) / "output.docx"
        doc = Document()
        doc.add_paragraph("测试正文。")
        add_page_number(doc, style="dash", position="outside")
        doc.save(source)

        custom = deepcopy(PRESETS["official"])
        custom.update({
            "name": "页码测试",
            "page_number_style": "page_text",
            "page_number_position": "center",
            "page_number_size": 12,
            "page_number_offset_mm": 10,
            "replace_existing_page_number": True,
        })
        format_document(str(source), str(output), preset_name="custom", custom_settings=custom)

        result = Document(output)
        section = result.sections[0]
        assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert "第 " in _footer_text(section.footer)
        assert " 页" in _footer_text(section.footer)
        assert abs(section.footer_distance.cm - 2.5) < 0.02

        with zipfile.ZipFile(output) as archive:
            footer_xml = "".join(
                archive.read(name).decode("utf-8")
                for name in archive.namelist()
                if name.startswith("word/footer") and name.endswith(".xml")
            )
        assert footer_xml.count("第 ") == 1
        assert "— " not in footer_xml


def test_non_page_footer_content_is_not_overwritten():
    doc = Document()
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "内部资料"

    add_page_number(doc, style="plain", replace_existing=True)

    assert footer.paragraphs[0].text == "内部资料"
    assert "PAGE" not in footer._element.xml


def test_format_document_resets_styles_and_strips_autospacing():
    with tempfile.TemporaryDirectory() as folder:
        source = Path(folder) / "source.docx"
        output = Path(folder) / "output.docx"
        doc = Document()
        doc.add_paragraph("一、第一项工作", style="Heading 3")
        doc.add_paragraph("正文段落。")
        doc.add_paragraph("（一）小标题", style="Heading 4")
        doc.save(source)

        format_document(str(source), str(output), preset_name="official")

        result = Document(output)
        for paragraph in result.paragraphs:
            if paragraph.text.strip():
                assert paragraph.style.name == "Normal"

        with zipfile.ZipFile(output) as archive:
            styles_xml = archive.read("word/styles.xml").decode("utf-8")
        assert "Autospacing" not in styles_xml
